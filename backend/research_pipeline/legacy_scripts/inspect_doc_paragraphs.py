import docx

doc = docx.Document("Crime severity scores.docx")
print("Total Paragraphs:", len(doc.paragraphs))
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if text:
        print(f"P{i}: {text}")
