import os
import docx
import pandas as pd

print("--- Inspecting 'Crime severity scores.docx' ---")
doc_path = "Crime severity scores.docx"
if os.path.exists(doc_path):
    doc = docx.Document(doc_path)
    print(f"Number of paragraphs: {len(doc.paragraphs)}")
    print(f"Number of tables: {len(doc.tables)}")
    # Print the first few paragraphs
    for i, p in enumerate(doc.paragraphs[:20]):
        if p.text.strip():
            print(f"Paragraph {i}: {p.text[:100]}")
    # Print table info if there are tables
    for i, table in enumerate(doc.tables):
        print(f"Table {i} has {len(table.rows)} rows and {len(table.columns)} columns.")
        # Print first row
        row_cells = [cell.text.strip() for cell in table.rows[0].cells]
        print(f"Row 0: {row_cells}")
        if len(table.rows) > 1:
            row1_cells = [cell.text.strip() for cell in table.rows[1].cells]
            print(f"Row 1: {row1_cells}")
else:
    print(f"'{doc_path}' not found!")

print("\n--- Inspecting 'Police force in England.xlsx' ---")
xlsx_path = "Police force in England.xlsx"
if os.path.exists(xlsx_path):
    xl = pd.ExcelFile(xlsx_path)
    print("Sheets in excel file:", xl.sheet_names)
    for sheet in xl.sheet_names:
        df = pd.read_excel(xlsx_path, sheet_name=sheet, nrows=5)
        print(f"\nSheet: {sheet}")
        print(df.head())
else:
    print(f"'{xlsx_path}' not found!")

print("\n--- Inspecting parquet files inside 'data' folder ---")
if os.path.exists("data"):
    files = os.listdir("data")
    parquet_files = [f for f in files if f.endswith(".parquet")]
    print("Parquet files found:", parquet_files)
    for pf in parquet_files:
        pf_path = os.path.join("data", pf)
        # Read just metadata or first row
        try:
            df = pd.read_parquet(pf_path)
            print(f"\nFile: {pf} (shape: {df.shape})")
            print("Columns:", df.columns.tolist())
            print(df.head(2))
        except Exception as e:
            print(f"Error reading {pf}: {e}")
else:
    print("'data' directory not found!")
